"""
Tests for docx_utils.html_to_docx_bytes — the universal HTML → Word
(.docx) converter used across all document modules so users can
upload their generated docs to Google Drive and edit them in Google
Docs without losing formatting.
"""
import io
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from docx import Document  # noqa: E402

from docx_utils import html_to_docx_bytes, NAVY, BLACK  # noqa: E402


SAMPLE_HTML = """
<h1>1. Executive Summary</h1>
<p>This is the <strong>introduction</strong> paragraph with <em>emphasis</em>.</p>
<h2>1.1 Key Findings</h2>
<ul>
  <li>First finding with <strong>bold</strong>.</li>
  <li>Second finding.</li>
</ul>
<h2>1.2 Data Table</h2>
<table>
  <thead>
    <tr><th>Metric</th><th>Value</th><th>Source</th></tr>
  </thead>
  <tbody>
    <tr><td>Small businesses</td><td>33.2 million</td><td>SBA (2023)</td></tr>
    <tr><td>Survival rate</td><td>50%</td><td>BLS (2023)</td></tr>
  </tbody>
</table>
<p>Closing paragraph after the table.</p>
"""


def _open(buf_bytes: bytes) -> Document:
    return Document(io.BytesIO(buf_bytes))


def test_basic_conversion_returns_valid_docx():
    out = html_to_docx_bytes(
        html=SAMPLE_HTML,
        title="Test Document",
        doc_type="Economic Impact Analysis",
        author="Leidis Pelaez",
        language="en",
    )
    assert isinstance(out, bytes)
    # Sanity: it must be openable by python-docx (i.e. valid OOXML).
    doc = _open(out)
    assert doc is not None


def test_cover_page_contains_metadata():
    out = html_to_docx_bytes(
        html=SAMPLE_HTML,
        title="Test Document",
        doc_type="Economic Impact Analysis",
        author="Leidis Pelaez",
        language="en",
        cover_subtitle="Prong 1 Analysis",
        legal_reference="Matter of Dhanasar, 26 I&N Dec. 884",
    )
    doc = _open(out)
    all_text = "\n".join(p.text for p in doc.paragraphs)
    assert "ECONOMIC IMPACT ANALYSIS" in all_text
    assert "Test Document" in all_text
    assert "Leidis Pelaez" in all_text
    assert "Project Proponent" in all_text
    assert "Prong 1 Analysis" in all_text
    assert "Matter of Dhanasar" in all_text


def test_body_headings_become_real_word_headings():
    out = html_to_docx_bytes(SAMPLE_HTML, title="Test", doc_type="Doc")
    doc = _open(out)
    # python-docx represents headings via the "Heading N" style.
    heading_texts = [
        p.text for p in doc.paragraphs
        if p.style and p.style.name and p.style.name.startswith("Heading")
    ]
    assert any("Executive Summary" in t for t in heading_texts)
    assert any("Key Findings" in t for t in heading_texts)
    assert any("Data Table" in t for t in heading_texts)


def test_tables_become_real_word_tables():
    out = html_to_docx_bytes(SAMPLE_HTML, title="Test", doc_type="Doc")
    doc = _open(out)
    assert len(doc.tables) == 1, f"Expected 1 table, got {len(doc.tables)}"
    tbl = doc.tables[0]
    # 3 rows × 3 cols (header + 2 data rows)
    assert len(tbl.rows) == 3
    assert len(tbl.columns) == 3
    header_cells = [c.text for c in tbl.rows[0].cells]
    assert "Metric" in header_cells
    assert "Value" in header_cells
    assert "Source" in header_cells
    # Spot check a data cell
    data = [c.text for c in tbl.rows[1].cells]
    assert "33.2 million" in data
    assert "SBA (2023)" in data


def test_lists_become_real_word_lists():
    out = html_to_docx_bytes(SAMPLE_HTML, title="Test", doc_type="Doc")
    doc = _open(out)
    bullet_paragraphs = [
        p.text for p in doc.paragraphs
        if p.style and p.style.name == "List Bullet"
    ]
    assert any("First finding" in t for t in bullet_paragraphs)
    assert any("Second finding" in t for t in bullet_paragraphs)


def test_inline_bold_italic_preserved():
    out = html_to_docx_bytes(SAMPLE_HTML, title="Test", doc_type="Doc")
    doc = _open(out)
    # find the introduction paragraph
    para = next(
        (p for p in doc.paragraphs if "introduction" in p.text and "emphasis" in p.text),
        None,
    )
    assert para is not None, "Could not find introduction paragraph"
    # Check that some run is bold (the "introduction" word) and another is italic.
    bolds = [r.text for r in para.runs if r.bold]
    italics = [r.text for r in para.runs if r.italic]
    assert any("introduction" in t for t in bolds), f"Bold runs: {bolds}"
    assert any("emphasis" in t for t in italics), f"Italic runs: {italics}"


def test_horizontal_rules_dropped():
    html = "<p>Before</p>\n<hr/>\n<p>After</p>\n---\n***\n<p>End</p>"
    out = html_to_docx_bytes(html, title="Test", doc_type="Doc", add_cover=False)
    doc = _open(out)
    body_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Before" in body_text
    assert "After" in body_text
    assert "End" in body_text
    # No `---` or `***` literal lines in body.
    for p in doc.paragraphs:
        stripped = p.text.strip()
        assert stripped not in ("---", "***", "___")


def test_no_cover_when_add_cover_false():
    out = html_to_docx_bytes(
        SAMPLE_HTML, title="Letter", doc_type="Letter",
        add_cover=False,
    )
    doc = _open(out)
    all_text = "\n".join(p.text for p in doc.paragraphs)
    # The doc_type banner ("LETTER") must NOT be at the top.
    assert not all_text.lstrip().startswith("LETTER")


def test_handles_subscript_and_superscript():
    html = "<p>The formula is M<sub>o</sub> = 1.85 with R<sup>2</sup> = 0.92.</p>"
    out = html_to_docx_bytes(html, title="Test", doc_type="Doc", add_cover=False)
    doc = _open(out)
    para = next((p for p in doc.paragraphs if "M" in p.text and "1.85" in p.text), None)
    assert para is not None
    sub_runs = [r for r in para.runs if r.font.subscript]
    sup_runs = [r for r in para.runs if r.font.superscript]
    assert any(r.text == 'o' for r in sub_runs)
    assert any(r.text == '2' for r in sup_runs)


def test_spanish_date_format():
    out = html_to_docx_bytes(
        "<p>Body</p>", title="Doc", doc_type="Reporte",
        author="X", language="es",
    )
    doc = _open(out)
    all_text = "\n".join(p.text for p in doc.paragraphs)
    # Should contain a Spanish month name.
    months = ["enero", "febrero", "marzo", "abril", "mayo", "junio",
              "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre"]
    assert any(m in all_text for m in months), (
        f"Expected a Spanish month in cover date. Body text: {all_text}"
    )
    assert "Proponente del Proyecto" in all_text


def test_handles_empty_html_safely():
    out = html_to_docx_bytes("", title="Test", doc_type="Doc")
    doc = _open(out)
    # Should not crash and should produce SOME content.
    assert len(doc.paragraphs) >= 1


# ── Patent (monochrome) export: headings must be BLACK, never navy ──────────

def test_mono_headings_are_black_not_navy():
    """Patents must be strictly black & white — heading runs render in black."""
    out = html_to_docx_bytes(
        html="<h2>FIELD OF THE INVENTION</h2><p>Body text.</p>",
        title="Patent", doc_type="Provisional Patent Application",
        language="en", add_cover=False, mono=True,
    )
    doc = _open(out)
    heading_runs = [
        run for p in doc.paragraphs if (p.style.name or "").startswith("Heading")
        for run in p.runs if run.text.strip()
    ]
    assert heading_runs, "no heading runs were produced"
    for run in heading_runs:
        col = run.font.color
        if col is not None and col.type is not None and col.rgb is not None:
            assert str(col.rgb) == str(BLACK), f"heading not black: {run.text!r} -> {col.rgb}"
            assert str(col.rgb) != str(NAVY)


def test_default_headings_stay_navy():
    """Non-mono documents keep the navy house colour (no regression)."""
    out = html_to_docx_bytes(
        html="<h2>Executive Summary</h2><p>Body.</p>",
        title="Plan", doc_type="Business Plan", language="en", add_cover=False,
    )
    doc = _open(out)
    colors = [
        str(run.font.color.rgb)
        for p in doc.paragraphs if (p.style.name or "").startswith("Heading")
        for run in p.runs
        if run.font.color is not None and run.font.color.type is not None and run.font.color.rgb is not None
    ]
    assert any(c == str(NAVY) for c in colors), f"expected navy headings, got {colors}"


def test_strong_renders_bold_without_markdown_artifacts():
    """<strong> must become real bold text — never literal '**' in the doc."""
    out = html_to_docx_bytes(
        html="<p>This is <strong>important</strong> and <em>noted</em>.</p>",
        title="Doc", doc_type="Doc", add_cover=False, mono=True,
    )
    doc = _open(out)
    body = "\n".join(p.text for p in doc.paragraphs)
    assert "**" not in body, f"literal markdown bold leaked: {body!r}"
    bold_runs = [r for p in doc.paragraphs for r in p.runs if r.bold and r.text.strip()]
    assert any("important" in r.text for r in bold_runs), "<strong> was not rendered bold"
