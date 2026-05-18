"""
docx_utils.py — Universal HTML → DOCX converter.

Used by every document module (policy papers, case studies, white papers,
econometric studies, patents, books, business plans, recommendation
letters, expert letters, intent letters, self-petition letters) to
produce a Microsoft Word .docx file from the SAME compiled HTML that
the PDF generators consume. The resulting .docx imports cleanly into
Google Docs (and Microsoft Word) preserving headings, paragraphs,
tables, lists, bold, italic, and basic styling — so users can edit
their generated documents without losing formatting.

Public API:
    html_to_docx_bytes(html, title, doc_type='document',
                       author=None, language='en',
                       project_title=None, applicant_name=None,
                       cover_title=None, cover_subtitle=None,
                       legal_reference=None) -> bytes

Why this is a separate module:
- Keeps server.py from growing further.
- Lets each router import a single helper without duplicating logic.
- Makes the conversion testable in isolation.

Design choices:
- We DO add a cover page (title + subtitle + author + date + optional
  legal reference) so the Word document feels as professional as the
  PDF.
- We do NOT attempt to reproduce decorative HRFlowable lines, banners,
  page numbers or background images — Word renders those poorly. The
  goal is clean, editable content.
- Inline HTML inside table cells is flattened to plain text + bold/italic
  (Word cell shading/grids are still applied).
"""
from __future__ import annotations

import io
import logging
import re
from datetime import datetime
from typing import Optional

from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Inches, Cm

logger = logging.getLogger(__name__)


# ── colour palette (matches PDF visual identity) ──────────────────────────
NAVY     = RGBColor(0x1A, 0x36, 0x5D)
GREY     = RGBColor(0x55, 0x55, 0x55)
LIGHT_GR = RGBColor(0x88, 0x88, 0x88)
HEADER_BG = "E2E8F0"   # light blue-grey for table header row


# ── helpers ───────────────────────────────────────────────────────────────
def _set_cell_shading(cell, hex_color: str) -> None:
    """Apply a background colour to a docx table cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)


def _set_cell_borders(cell) -> None:
    """Apply thin borders to a docx table cell (matches the PDF grid look)."""
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:color'), 'A0AEC0')
        tc_borders.append(b)
    tc_pr.append(tc_borders)


def _add_page_break(doc: Document) -> None:
    """Insert an explicit page break paragraph."""
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


def _add_runs_from_inline_html(paragraph, node) -> None:
    """Walk an inline HTML element/string and add runs to a docx paragraph
    preserving <strong>/<b>, <em>/<i>, <sub>, <sup>, <code>, <br/>."""
    if isinstance(node, NavigableString):
        text = str(node)
        if text:
            paragraph.add_run(text)
        return

    if not isinstance(node, Tag):
        return

    name = node.name.lower() if node.name else ''

    if name == 'br':
        paragraph.add_run().add_break()
        return

    # Compute inherited formatting flags from this element's own tag.
    bold = name in ('strong', 'b')
    italic = name in ('em', 'i')
    sub = name == 'sub'
    sup = name == 'sup'
    code = name == 'code'

    if not node.children:
        text = node.get_text() or ''
        if text:
            run = paragraph.add_run(text)
            if bold:    run.bold = True
            if italic:  run.italic = True
            if sub:     run.font.subscript = True
            if sup:     run.font.superscript = True
            if code:    run.font.name = 'Consolas'
        return

    for child in node.children:
        if isinstance(child, NavigableString):
            text = str(child)
            if text:
                run = paragraph.add_run(text)
                if bold:    run.bold = True
                if italic:  run.italic = True
                if sub:     run.font.subscript = True
                if sup:     run.font.superscript = True
                if code:    run.font.name = 'Consolas'
        elif isinstance(child, Tag):
            child_name = (child.name or '').lower()
            if child_name == 'br':
                paragraph.add_run().add_break()
                continue
            # Inline-style children — recurse and combine flags.
            child_bold   = bold   or child_name in ('strong', 'b')
            child_italic = italic or child_name in ('em', 'i')
            child_sub    = sub    or child_name == 'sub'
            child_sup    = sup    or child_name == 'sup'
            child_code   = code   or child_name == 'code'
            for sub_node in child.children:
                if isinstance(sub_node, NavigableString):
                    text = str(sub_node)
                    if text:
                        run = paragraph.add_run(text)
                        if child_bold:    run.bold = True
                        if child_italic:  run.italic = True
                        if child_sub:     run.font.subscript = True
                        if child_sup:     run.font.superscript = True
                        if child_code:    run.font.name = 'Consolas'
                elif isinstance(sub_node, Tag):
                    # Two-level nested inline (e.g. <strong><em>x</em></strong>) — recurse.
                    _add_runs_from_inline_html(paragraph, sub_node)


def _process_block(doc: Document, element: Tag) -> None:
    """Convert one block-level HTML element into a docx flowable."""
    name = (element.name or '').lower()

    # ── Headings ────────────────────────────────────────────────────────
    if name in ('h1', 'h2', 'h3', 'h4', 'h5', 'h6'):
        level = int(name[1])
        # docx supports built-in Heading 1..9
        h = doc.add_heading(level=min(level, 4))
        # Clear the default run created by add_heading (empty)
        for run in list(h.runs):
            run.text = ''
        _add_runs_from_inline_html(h, element)
        # Override colour to navy for visual consistency with PDFs.
        for run in h.runs:
            run.font.color.rgb = NAVY
        return

    # ── Tables ──────────────────────────────────────────────────────────
    if name == 'table':
        _emit_table(doc, element)
        return

    # ── Lists ───────────────────────────────────────────────────────────
    if name == 'ul':
        for li in element.find_all('li', recursive=False):
            p = doc.add_paragraph(style='List Bullet')
            _add_runs_from_inline_html(p, li)
        return
    if name == 'ol':
        for li in element.find_all('li', recursive=False):
            p = doc.add_paragraph(style='List Number')
            _add_runs_from_inline_html(p, li)
        return

    # ── Horizontal rules → drop entirely (PDF policy) ───────────────────
    if name == 'hr':
        return

    # ── Block quote ─────────────────────────────────────────────────────
    if name == 'blockquote':
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.4)
        _add_runs_from_inline_html(p, element)
        return

    # ── Pre / code block ────────────────────────────────────────────────
    if name == 'pre':
        text = element.get_text()
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = 'Consolas'
        run.font.size = Pt(10)
        return

    # ── div: recurse into children (treat as transparent wrapper) ───────
    if name == 'div':
        for child in element.children:
            if isinstance(child, Tag):
                _process_block(doc, child)
            elif isinstance(child, NavigableString):
                text = str(child).strip()
                if text:
                    p = doc.add_paragraph()
                    p.add_run(text)
        return

    # ── img: skip (Word handles inline imgs poorly without size hints) ──
    if name == 'img':
        # Try to add as inline if it's a base64 data URI; otherwise skip.
        src = element.get('src', '')
        if src.startswith('data:image'):
            try:
                import base64
                if 'base64,' in src:
                    b64 = src.split('base64,', 1)[1]
                    img_bytes = base64.b64decode(b64)
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run()
                    run.add_picture(io.BytesIO(img_bytes), width=Inches(5.5))
                    return
            except Exception as exc:
                logger.warning("Skipping inline image: %s", exc)
        # else: omit the image
        return

    # ── Default: paragraph ──────────────────────────────────────────────
    if name == 'p' or name == '':
        text_only = element.get_text(strip=True) if hasattr(element, 'get_text') else ''
        if not text_only:
            return
        p = doc.add_paragraph()
        _add_runs_from_inline_html(p, element)
        return

    # Unknown tag → fall through as a paragraph.
    p = doc.add_paragraph()
    _add_runs_from_inline_html(p, element)


def _emit_table(doc: Document, table_tag: Tag) -> None:
    """Convert a <table> into a real docx Table with header shading and borders."""
    # Collect rows from <thead> and <tbody> (and bare <tr>).
    rows: list[list[str]] = []
    is_header_row: list[bool] = []

    thead = table_tag.find('thead')
    if thead:
        for tr in thead.find_all('tr', recursive=False):
            cells = [c for c in tr.find_all(['th', 'td'], recursive=False)]
            rows.append(cells)
            is_header_row.append(True)

    tbody = table_tag.find('tbody')
    if tbody:
        for tr in tbody.find_all('tr', recursive=False):
            cells = [c for c in tr.find_all(['th', 'td'], recursive=False)]
            rows.append(cells)
            is_header_row.append(False)
    else:
        # bare <tr> children of <table>
        for tr in table_tag.find_all('tr', recursive=False):
            if thead and tr.parent is thead:
                continue
            cells = [c for c in tr.find_all(['th', 'td'], recursive=False)]
            if cells:
                rows.append(cells)
                is_header_row.append(False)

    if not rows:
        return

    n_cols = max(len(r) for r in rows)
    if n_cols == 0:
        return

    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.style = 'Table Grid'
    table.autofit = True

    for r_idx, (cell_tags, is_hdr) in enumerate(zip(rows, is_header_row)):
        for c_idx in range(n_cols):
            cell = table.rows[r_idx].cells[c_idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            # Clear the default empty paragraph
            cell.text = ''
            paragraph = cell.paragraphs[0]
            if c_idx < len(cell_tags):
                _add_runs_from_inline_html(paragraph, cell_tags[c_idx])
            if is_hdr:
                _set_cell_shading(cell, HEADER_BG)
                # Bold + navy header text
                for run in paragraph.runs:
                    run.bold = True
                    run.font.color.rgb = NAVY
            _set_cell_borders(cell)
    # Trailing spacer paragraph after the table
    doc.add_paragraph()


# ── public API ────────────────────────────────────────────────────────────
def html_to_docx_bytes(
    html: str,
    title: str,
    doc_type: str = 'document',
    author: Optional[str] = None,
    language: str = 'en',
    cover_subtitle: Optional[str] = None,
    legal_reference: Optional[str] = None,
    add_cover: bool = True,
) -> bytes:
    """
    Convert HTML to .docx bytes.

    Parameters
    ----------
    html : str
        The same compiled HTML the PDF generators receive.
    title : str
        Document title (also used in metadata + cover).
    doc_type : str
        Free-text type label shown on the cover (e.g. "Economic Impact
        Analysis", "Case Study", "White Paper").
    author : Optional[str]
        Shown as "Author / Project Proponent" on the cover.
    language : str
        'en' or 'es' — affects date formatting and labels on the cover.
    cover_subtitle : Optional[str]
        Optional subtitle shown under the title on the cover.
    legal_reference : Optional[str]
        Optional one-liner shown at the bottom of the cover (e.g.
        "Prepared pursuant to Matter of Dhanasar, 26 I&N Dec. 884
        (AAO 2016)").
    add_cover : bool
        If False, skip the cover page entirely (useful for short letters).
    """
    doc = Document()

    # ── document defaults: serif body font, 11pt, 1-inch margins ────────
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # ── document properties (visible in Word's File ▸ Info ▸ Properties)
    doc.core_properties.title = title or 'Document'
    if author:
        doc.core_properties.author = author

    # ── COVER PAGE ──────────────────────────────────────────────────────
    if add_cover:
        # Vertical breathing room
        for _ in range(6):
            doc.add_paragraph()

        # Banner — doc_type ALL CAPS in navy
        banner = doc.add_paragraph()
        banner.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = banner.add_run((doc_type or 'Document').upper())
        run.bold = True
        run.font.size = Pt(22)
        run.font.color.rgb = NAVY

        # Title
        if title:
            title_p = doc.add_paragraph()
            title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = title_p.add_run(title)
            run.bold = True
            run.font.size = Pt(16)
            run.font.color.rgb = NAVY

        # Author / Project Proponent
        if author:
            label = "Project Proponent" if language == 'en' else "Proponente del Proyecto"
            ap = doc.add_paragraph()
            ap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = ap.add_run(f"{label}: {author}")
            run.font.size = Pt(12)
            run.font.color.rgb = GREY

        # Subtitle
        if cover_subtitle:
            for _ in range(2):
                doc.add_paragraph()
            sp = doc.add_paragraph()
            sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = sp.add_run(cover_subtitle)
            run.bold = True
            run.font.size = Pt(13)

        # Date
        for _ in range(3):
            doc.add_paragraph()
        dp = doc.add_paragraph()
        dp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if language == 'es':
            spanish_months = {
                1: "enero", 2: "febrero", 3: "marzo", 4: "abril",
                5: "mayo", 6: "junio", 7: "julio", 8: "agosto",
                9: "septiembre", 10: "octubre", 11: "noviembre", 12: "diciembre",
            }
            now = datetime.now()
            date_text = f"{now.day} de {spanish_months[now.month]} de {now.year}"
        else:
            date_text = datetime.now().strftime("%B %d, %Y")
        run = dp.add_run(date_text)
        run.font.size = Pt(11)
        run.font.color.rgb = GREY

        # Legal reference (optional)
        if legal_reference:
            doc.add_paragraph()
            lp = doc.add_paragraph()
            lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = lp.add_run(legal_reference)
            run.italic = True
            run.font.size = Pt(9)
            run.font.color.rgb = LIGHT_GR

        _add_page_break(doc)

    # ── BODY ────────────────────────────────────────────────────────────
    body_html = (html or '').strip()
    if not body_html:
        doc.add_paragraph("(No content)")
    else:
        # Defensive cleanups: drop residual markdown HR markers, the
        # markdown-table separator rows that survived as text, and stray
        # stand-alone pipe rows.
        body_html = re.sub(r'(?m)^\s*[-*_=]{3,}\s*$', '', body_html)

        soup = BeautifulSoup(body_html, 'html.parser')
        # If the soup has no block-level children (rare), wrap into <div>.
        top_level_blocks = [c for c in soup.children if isinstance(c, Tag)]
        if top_level_blocks:
            for child in top_level_blocks:
                _process_block(doc, child)
            # Also grab any trailing NavigableString text
            for child in soup.children:
                if isinstance(child, NavigableString):
                    text = str(child).strip()
                    if text:
                        doc.add_paragraph(text)
        else:
            # Plain text only
            for line in body_html.split('\n'):
                line = line.strip()
                if line:
                    doc.add_paragraph(line)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


__all__ = ["html_to_docx_bytes", "build_docx_response", "DOCX_MIME"]


# ── FastAPI helper ────────────────────────────────────────────────────────
DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def build_docx_response(
    *,
    content: str,
    title: str,
    filename_stem: str,
    doc_type: str,
    author: Optional[str] = None,
    language: str = 'en',
    cover_subtitle: Optional[str] = None,
    legal_reference: Optional[str] = None,
    add_cover: bool = True,
    is_html: bool = False,  # accepted for back-compat; ignored — see notes
):
    """
    Convenience wrapper used by every /download-docx endpoint.

    Internally we run `markdown.markdown()` on the input so that:
      - Pure-markdown content gets fully converted.
      - Mixed content (HTML wrappers + bare markdown tables, which is how
        Claude Opus often emits econometric-study sections) ALSO gets its
        markdown features expanded.
      - Pre-broken tables (`<p>| ... |</p>`) are unwrapped first so the
        `tables` extension can recognize them.
      - Pure HTML passes through markdown.markdown() as a no-op.

    Returns a `fastapi.responses.StreamingResponse` ready to be returned
    from your route handler.
    """
    from fastapi.responses import StreamingResponse  # local import to avoid hard dep on top
    import markdown as _md

    raw = content or ''

    # 1. Unwrap pre-broken pipe-rows that got wrapped in <p> by an earlier
    #    HTML conversion (very common in econometric-study DB content).
    #    Convert <strong>/<em> inside cells to ** / * markers FIRST so
    #    the unwrap regex (which forbids `<` inside cells) still matches.
    raw = re.sub(r'<strong>\s*(.+?)\s*</strong>', r'**\1**', raw, flags=re.DOTALL)
    raw = re.sub(r'<em>\s*(.+?)\s*</em>', r'*\1*', raw, flags=re.DOTALL)
    raw = re.sub(r'<p[^>]*>\s*(\|[^<\n]+\|)\s*</p>', r'\1', raw)

    # 2. Inject blank line BEFORE/AFTER each pipe-table block so
    #    `tables` extension recognises them.
    _lines = raw.split('\n')
    _fixed = []
    _prev_was_pipe = False
    for _i, _ln in enumerate(_lines):
        _stripped = _ln.strip()
        _is_pipe = _stripped.startswith('|') and _stripped.endswith('|')
        if _is_pipe and not _prev_was_pipe and _i > 0:
            _prev = _lines[_i - 1].strip()
            if _prev != '':
                _fixed.append('')
        if _prev_was_pipe and not _is_pipe and _stripped != '':
            _fixed.append('')
        _fixed.append(_ln)
        _prev_was_pipe = _is_pipe
    raw = '\n'.join(_fixed)

    # 3. Auto-inject separator rows for 2-row tables that lack one
    #    (Claude's 2-column key/value tables tend to skip the |---|---|).
    _lines = raw.split('\n')
    _out = []
    _i = 0
    while _i < len(_lines):
        _s = _lines[_i].strip()
        if _s.startswith('|') and _s.endswith('|'):
            _block_start = _i
            while _i < len(_lines):
                _ss = _lines[_i].strip()
                if _ss.startswith('|') and _ss.endswith('|'):
                    _i += 1
                else:
                    break
            _block = _lines[_block_start:_i]
            _has_sep = any(re.match(r'^\s*\|[\s\-:|]+\|\s*$', _l) and _l.count('-') >= 2 for _l in _block)
            _out.append(_block[0])
            if not _has_sep and _block:
                _cols = _block[0].count('|') - 1
                if _cols >= 1:
                    _out.append('|' + '|'.join(['---'] * _cols) + '|')
            _out.extend(_block[1:])
        else:
            _out.append(_lines[_i])
            _i += 1
    raw = '\n'.join(_out)

    # 4. Run markdown → HTML (no-op for pure HTML, expands markdown otherwise).
    html = _md.markdown(raw, extensions=['tables', 'fenced_code', 'nl2br'])

    docx_bytes = html_to_docx_bytes(
        html=html,
        title=title,
        doc_type=doc_type,
        author=author,
        language=language,
        cover_subtitle=cover_subtitle,
        legal_reference=legal_reference,
        add_cover=add_cover,
    )

    safe_stem = re.sub(r'[^A-Za-z0-9_\-]', '_', filename_stem)[:80] or "document"
    filename = f"{safe_stem}_{language.upper()}.docx"
    return StreamingResponse(
        io.BytesIO(docx_bytes),
        media_type=DOCX_MIME,
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )
