"""
Robust PDF text extraction with multi-engine fallback and Vision OCR rescue.

Used by ALL letter generation flows (expert, recommendation, intent) and any
other module that needs to read user-uploaded CVs / project docs.

Why this module exists (historical bug 2026-05-28):
  - Letter flows were using PyPDF2 alone to read uploaded PDFs.
  - CVs designed in Canva / Adobe Express ship the text as vector paths (no
    extractable text layer) so PyPDF2 returned an empty string.
  - The downstream LLM extractor then HALLUCINATED a generic signer
    ("Dr. Emily Thompson, MIT") because it had nothing real to anchor on.
  - Fix: cascade of 5 PDF engines (PyMuPDF → pdfplumber → pypdfium2 → PyPDF2
    → pdfminer.six) plus GPT-4o Vision OCR as last resort for vectorized PDFs.

Also exposes `validate_signer_name_in_cv()` which the letter flows call to
catch any hallucinated signer that slipped past the prompt.
"""
import base64
import hashlib
import io
import logging
import re
import time
from collections import OrderedDict
from typing import Awaitable, Callable, Optional

logger = logging.getLogger(__name__)


# ──────────────────────────────────────────────────────────────────────────────
# Cache de extracciones por hash SHA-256 del contenido del PDF.
#
# Caso de uso real: una operadora hace 30 cartas en 4 min subiendo el MISMO
# CV del cliente y el MISMO project_info en cada carta (solo el firmante
# cambia). Sin cache: 90 extracciones, ~30 con OCR de 15s = 7.5min solo OCR.
# Con cache: ~32 extracciones reales (1 cliente + 1 proyecto + 30 firmantes).
#
# LRU + TTL en memoria del proceso. No se persiste a disco para mantenerlo
# simple — si reiniciás el container se pierde el cache, pero es trivial
# regenerarlo en pocas extracciones.
# ──────────────────────────────────────────────────────────────────────────────

_EXTRACT_CACHE: "OrderedDict[str, tuple[float, str]]" = OrderedDict()
_EXTRACT_CACHE_MAX = 256        # entradas — cada CV ocupa ~10-50KB en RAM
_EXTRACT_CACHE_TTL = 3600.0     # 1 hora — alcanza para una sesión de operadora


def _cache_key(content: bytes) -> str:
    """SHA-256 truncado a 16 chars — colisión astronómicamente improbable
    para los volúmenes que manejamos (<10K PDFs/día)."""
    return hashlib.sha256(content).hexdigest()[:16]


def _cache_get(key: str) -> Optional[str]:
    entry = _EXTRACT_CACHE.get(key)
    if entry is None:
        return None
    ts, text = entry
    if time.time() - ts > _EXTRACT_CACHE_TTL:
        _EXTRACT_CACHE.pop(key, None)
        return None
    # Refrescar posición LRU (más reciente al final).
    _EXTRACT_CACHE.move_to_end(key)
    return text


def _cache_put(key: str, text: str) -> None:
    _EXTRACT_CACHE[key] = (time.time(), text)
    _EXTRACT_CACHE.move_to_end(key)
    while len(_EXTRACT_CACHE) > _EXTRACT_CACHE_MAX:
        _EXTRACT_CACHE.popitem(last=False)  # drop LRU


# ──────────────────────────────────────────────────────────────────────────────
# Individual extraction engines — each tries its best and may raise on failure.
# Order in the cascade reflects observed quality on real-world CVs.
# ──────────────────────────────────────────────────────────────────────────────

def _extract_pdf_with_fitz(content: bytes) -> str:
    """PyMuPDF — most robust for digital PDFs with complex layouts."""
    import fitz  # PyMuPDF
    doc = fitz.open(stream=content, filetype="pdf")
    try:
        return "\n".join(page.get_text() or "" for page in doc)
    finally:
        doc.close()


def _extract_pdf_with_pdfplumber(content: bytes) -> str:
    import pdfplumber
    with pdfplumber.open(io.BytesIO(content)) as pdf:
        return "\n".join(page.extract_text() or "" for page in pdf.pages)


def _extract_pdf_with_pypdfium2(content: bytes) -> str:
    import pypdfium2 as pdfium
    doc = pdfium.PdfDocument(content)
    try:
        out = []
        for page in doc:
            tp = page.get_textpage()
            try:
                out.append(tp.get_text_range() or "")
            finally:
                tp.close()
            page.close()
        return "\n".join(out)
    finally:
        doc.close()


def _extract_pdf_with_pypdf2(content: bytes) -> str:
    from PyPDF2 import PdfReader
    reader = PdfReader(io.BytesIO(content))
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def _extract_pdf_with_pdfminer(content: bytes) -> str:
    from pdfminer.high_level import extract_text as _miner_extract
    return _miner_extract(io.BytesIO(content)) or ""


def _tesseract_one_page(pil_image, *, page_idx: int, n_pages: int, file_label: str) -> str:
    """OCR a single PIL image with Tesseract — runs in a thread pool worker."""
    import pytesseract  # type: ignore
    try:
        page_text = pytesseract.image_to_string(pil_image, lang="eng+spa") or ""
    except pytesseract.TesseractNotFoundError:
        # Re-raise so the caller falls back to the next OCR engine.
        raise
    except Exception as e:
        logger.warning(
            f"🔍 Tesseract {file_label} página {page_idx + 1} falló: {e!r}"
        )
        return ""
    logger.error(
        f"🔍 Tesseract {file_label} página {page_idx + 1}/{n_pages} → "
        f"{len(page_text)}c"
    )
    return page_text.strip()


# Semaphore global: limita cuántas páginas de Tesseract OCR corren al mismo
# tiempo POR PROCESO (afecta todas las cartas concurrentes).
#
# Tesseract es CPU-bound: cada página usa ~1 core a 100%. Si tenemos N cores
# disponibles, dejamos 1 libre para el event loop async + I/O y usamos N-1
# para OCR. Cap a 8 para no abrir demasiados workers en máquinas grandes.
#
# Ejemplos:
#   container con 2 vCPUs → 2 slots (Tesseract a 100%, async un poco lento pero ok)
#   container con 4 vCPUs → 3 slots
#   container con 8 vCPUs → 7 slots
#   container con 16+ vCPUs → 8 slots (cap)
#
# Con N cartas concurrentes pidiendo OCR, las páginas se serializan a través
# de este semaphore; no hay timeouts en la cola, solo en la ejecución (45s/pág
# + 180s total por CV). Si hay más cartas que slots, las demás esperan turno.
_TESSERACT_SEMAPHORE: Optional["object"] = None  # asyncio.Semaphore creado lazy
_TESSERACT_SLOTS: int = 0  # se setea al crear el semaphore


def _get_tesseract_semaphore():
    """Lazy init del semaphore — solo se crea cuando hay event loop activo.

    Cada worker uvicorn tiene su PROPIO proceso → su propio semaphore. Con
    4 workers × ~3 slots c/u = 12 OCRs en paralelo en todo el container,
    suficiente para picos de 30 cartas en 4 minutos (donde además el cache
    LRU short-circuitea los PDFs repetidos).

    El cap por proceso se puede tunear con env var TESSERACT_MAX_CONCURRENCY.
    """
    import asyncio
    import os
    global _TESSERACT_SEMAPHORE, _TESSERACT_SLOTS
    if _TESSERACT_SEMAPHORE is None:
        cpu_count = os.cpu_count() or 2
        cap = int(os.environ.get("TESSERACT_MAX_CONCURRENCY", "6"))
        # Dejar 1 core para event loop. Min 2, max `cap`.
        _TESSERACT_SLOTS = max(2, min(cpu_count - 1, cap))
        _TESSERACT_SEMAPHORE = asyncio.Semaphore(_TESSERACT_SLOTS)
        logger.error(
            f"🔧 Tesseract OCR concurrency: {_TESSERACT_SLOTS} slots/worker "
            f"(cpu_count={cpu_count}, cap={cap}). Multi-worker uvicorn "
            f"multiplica esto por #workers."
        )
    return _TESSERACT_SEMAPHORE


async def _ocr_pdf_with_tesseract(
    content: bytes,
    *,
    max_pages: int = 8,
    file_label: str = "",
    per_page_timeout: float = 45.0,
    total_timeout: float = 180.0,
) -> str:
    """Local OCR via Tesseract — gratis, sin red, sin depender de OpenAI.

    Procesa páginas con paralelismo LIMITADO (max 2 simultáneas via semaphore)
    para no saturar el CPU del container. Cada página tiene su propio
    timeout (`per_page_timeout`), y la operación completa tiene un techo
    global (`total_timeout`) para que nunca se cuelgue indefinidamente.

    Funciona para PDFs vectorizados (Canva / Adobe Express) y escaneados.
    Requiere `tesseract-ocr` en el sistema (Dockerfile) + `pytesseract`.

    Configurado para Inglés + Español (CVs URPE típicamente mezclan ambos).
    """
    import asyncio
    import pypdfium2 as pdfium

    doc = pdfium.PdfDocument(content)
    pages_pil = []
    try:
        n_pages = min(len(doc), max_pages)
        if len(doc) > max_pages:
            logger.error(
                f"🔍 Tesseract {file_label}: PDF tiene {len(doc)} páginas, "
                f"limitando OCR a las primeras {max_pages}"
            )
        # scale=2.0 ≈ 144 DPI — buen balance precisión/velocidad.
        for i in range(n_pages):
            page = doc[i]
            pages_pil.append(page.render(scale=2.0).to_pil())
            page.close()
    finally:
        doc.close()

    if not pages_pil:
        return ""

    logger.error(
        f"🔍 Tesseract {file_label}: renderizadas {len(pages_pil)} páginas, "
        f"iniciando OCR (max 2 simultáneas, timeout {per_page_timeout}s/pág)..."
    )

    sem = _get_tesseract_semaphore()
    loop = asyncio.get_event_loop()

    async def _run_one(idx: int, img):
        async with sem:
            try:
                # Cada página tiene su propio timeout. Si una página se cuelga
                # (PDF corrupto, página rara), no detiene a las demás.
                return await asyncio.wait_for(
                    loop.run_in_executor(
                        None,
                        lambda: _tesseract_one_page(
                            img, page_idx=idx, n_pages=len(pages_pil),
                            file_label=file_label,
                        ),
                    ),
                    timeout=per_page_timeout,
                )
            except asyncio.TimeoutError:
                logger.error(
                    f"🔍 Tesseract {file_label} página {idx + 1} TIMEOUT "
                    f"({per_page_timeout}s) — skip"
                )
                return ""

    tasks = [_run_one(idx, img) for idx, img in enumerate(pages_pil)]

    try:
        # Timeout global para que el OCR completo nunca pase de `total_timeout`.
        results = await asyncio.wait_for(
            asyncio.gather(*tasks, return_exceptions=True),
            timeout=total_timeout,
        )
    except asyncio.TimeoutError:
        logger.error(
            f"🔍 Tesseract {file_label}: TIMEOUT GLOBAL ({total_timeout}s) — "
            f"abortando OCR completo"
        )
        return ""

    text_parts = []
    for r in results:
        if isinstance(r, Exception):
            if r.__class__.__name__ == "TesseractNotFoundError":
                logger.error(
                    "🔍 Tesseract binary NOT installed in container. "
                    "Add `tesseract-ocr tesseract-ocr-eng tesseract-ocr-spa` "
                    "to the Dockerfile apt-get install. Skipping Tesseract OCR."
                )
                raise r
            logger.warning(f"🔍 Tesseract {file_label} task falló: {r!r}")
            continue
        if r:
            text_parts.append(r)
    return "\n\n".join(text_parts)


# ──────────────────────────────────────────────────────────────────────────────
# Vision OCR rescue — only fires when ALL 5 engines above returned empty/junk.
# Needed for PDFs whose text was converted to vector paths (Canva / Adobe
# Express exports) or scanned PDFs. Costs ~$0.02-0.08 per CV via GPT-4o.
# ──────────────────────────────────────────────────────────────────────────────

VisionFn = Callable[..., Awaitable[str]]


async def _ocr_pdf_with_vision(
    content: bytes,
    vision_fn: VisionFn,
    *,
    max_pages: int = 8,
    file_label: str = "",
    per_page_timeout: float = 45.0,
) -> str:
    """OCR each page via GPT-4o Vision (or whatever `vision_fn` plugs in).

    `vision_fn` must be an async callable accepting (messages, max_tokens,
    temperature) and returning the model's text response — same signature as
    `server.call_openai_gpt4o_vision`.
    """
    import pypdfium2 as pdfium
    from PIL import Image  # type: ignore

    doc = pdfium.PdfDocument(content)
    try:
        n_pages = min(len(doc), max_pages)
        if len(doc) > max_pages:
            logger.warning(
                f"🔍 {file_label} tiene {len(doc)} páginas, OCR limitado a las primeras {max_pages}"
            )
        all_text_parts = []
        for i in range(n_pages):
            page = doc[i]
            pil = page.render(scale=2.0).to_pil()  # ~144 DPI
            page.close()
            w, h = pil.size
            longest = max(w, h)
            if longest > 2000:
                ratio = 2000 / longest
                pil = pil.resize((int(w * ratio), int(h * ratio)), Image.LANCZOS)
            buf = io.BytesIO()
            pil.convert("RGB").save(buf, format="JPEG", quality=85)
            jpeg_b64 = base64.b64encode(buf.getvalue()).decode("ascii")
            data_url = f"data:image/jpeg;base64,{jpeg_b64}"

            messages = [
                {
                    "role": "system",
                    "content": (
                        "You are an OCR engine. Extract the literal visible text from "
                        "the image in reading order. Preserve line breaks. Do NOT "
                        "translate, summarize, paraphrase, comment, or add anything "
                        "that is not visible on the page. If the page is blank or "
                        "unreadable, respond with an empty string."
                    ),
                },
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "text",
                            "text": (
                                "Transcribe all text visible in this page exactly as "
                                "it appears. Return only the text — no extra commentary."
                            ),
                        },
                        {"type": "image_url", "image_url": {"url": data_url}},
                    ],
                },
            ]
            try:
                import asyncio as _asyncio
                page_text = await _asyncio.wait_for(
                    vision_fn(messages=messages, max_tokens=2500, temperature=0.0),
                    timeout=per_page_timeout,
                )
                if page_text:
                    all_text_parts.append(page_text.strip())
                logger.error(
                    f"🔍 Vision OCR {file_label} página {i + 1}/{n_pages} → "
                    f"{len(page_text or '')}c"
                )
            except _asyncio.TimeoutError:
                logger.error(
                    f"🔍 Vision OCR {file_label} página {i + 1} TIMEOUT "
                    f"({per_page_timeout}s) — skip"
                )
                continue
            except Exception as e:
                logger.warning(f"🔍 OCR {file_label} página {i + 1} falló: {e!r}")
                continue
        return "\n\n".join(all_text_parts)
    finally:
        doc.close()


# ──────────────────────────────────────────────────────────────────────────────
# Public API
# ──────────────────────────────────────────────────────────────────────────────

def _looks_like_pdf(content: bytes, filename: str, content_type: str) -> bool:
    """PDF detection that doesn't trust the filename alone.

    A browser can send a PDF with filename=``cv`` (no extension) or
    content_type=``application/octet-stream``. The %PDF- magic header is the
    authoritative signal.
    """
    if content.startswith(b'%PDF-'):
        return True
    if filename.lower().endswith('.pdf'):
        return True
    if 'pdf' in (content_type or '').lower():
        return True
    return False


async def extract_text_from_bytes(
    content: bytes,
    filename: str,
    *,
    vision_fn: Optional[VisionFn] = None,
    content_type: str = "",
    file_label: str = "",
) -> str:
    """Extract text from file bytes — robust 6-engine cascade for PDFs.

    Args:
        content: raw bytes of the uploaded file.
        filename: original filename (used for extension detection / logging).
        vision_fn: optional async callable for OCR rescue. If None, OCR is
            skipped and the empty string is returned when all 5 PDF engines
            fail. Letter flows should always pass one in.
        content_type: MIME type if available.
        file_label: short human-readable label for log lines (e.g. "expert_cv").

    Returns:
        Extracted text. Empty string if everything failed.
    """
    fname_raw = filename or ""
    fname = fname_raw.lower()
    label = file_label or fname_raw or "<unnamed>"
    first_bytes = content[:8]
    logger.error(
        f"📥 extract_text | label={label!r} filename={fname_raw!r} "
        f"content_type={content_type!r} bytes={len(content)} magic={first_bytes!r}"
    )

    # ── Cache hit? — short-circuit ANTES de tocar Tesseract o Vision OCR ──────
    # Crítico cuando una operadora hace 30 cartas en serie con el mismo CV
    # de cliente / project_info — ahorra ~30 OCRs redundantes.
    cache_key = _cache_key(content)
    cached = _cache_get(cache_key)
    if cached is not None:
        logger.error(
            f"📥 {label}: ✅ CACHE HIT (sha256={cache_key}) — devolviendo "
            f"{len(cached)}c sin re-extraer"
        )
        return cached

    if _looks_like_pdf(content, fname, content_type):
        if not fname.endswith('.pdf'):
            logger.warning(
                f"📥 {label}: tratando como PDF por magic/content-type "
                f"aunque filename no termina en .pdf | filename={fname_raw!r}"
            )
        extractors = [
            ("PyMuPDF", _extract_pdf_with_fitz),
            ("pdfplumber", _extract_pdf_with_pdfplumber),
            ("pypdfium2", _extract_pdf_with_pypdfium2),
            ("PyPDF2", _extract_pdf_with_pypdf2),
            ("pdfminer", _extract_pdf_with_pdfminer),
        ]
        last_err = None
        best_text = ""
        for name, fn in extractors:
            try:
                text = fn(content) or ""
                non_ws = len(re.sub(r'\s+', '', text))
                if non_ws >= 50:
                    logger.info(
                        f"📄 {label}: PDF extraído con {name} | "
                        f"{len(text)}c totales, {non_ws}c no-whitespace"
                    )
                    _cache_put(cache_key, text)
                    return text
                if non_ws > len(re.sub(r'\s+', '', best_text)):
                    best_text = text
                logger.warning(
                    f"📄 {label}: {name} extrajo muy poco texto ({non_ws}c) — probando siguiente"
                )
            except Exception as e:
                last_err = e
                logger.warning(f"📄 {label}: {name} falló: {e!r} — probando siguiente")

        logger.error(
            f"📄 {label}: todos los 5 extractores PDF dieron poco texto. "
            f"Mejor intento: {len(best_text)}c. Último error: {last_err!r}. "
            f"Iniciando cascada de OCR..."
        )

        # ── Escalón 6: Tesseract OCR (local, gratis, sin red) ─────────────────
        # Usamos ERROR como nivel para que SIEMPRE sea visible en logs.
        # Tesseract maneja bien PDFs vectorizados y escaneados sin costo, y
        # procesa todas las páginas en paralelo para no bloquear el background.
        try:
            logger.error(f"🔍 {label}: intentando Tesseract OCR (local, paralelo)...")
            tess_text = await _ocr_pdf_with_tesseract(content, file_label=label)
            non_ws_tess = len(re.sub(r'\s+', '', tess_text or ''))
            if non_ws_tess >= 50:
                logger.error(
                    f"🔍 {label}: ✅ Tesseract OCR rescató el PDF | "
                    f"{len(tess_text)}c totales, {non_ws_tess}c no-whitespace"
                )
                _cache_put(cache_key, tess_text)
                return tess_text
            logger.error(
                f"🔍 {label}: Tesseract OCR también dio poco texto ({non_ws_tess}c)"
            )
            if non_ws_tess > len(re.sub(r'\s+', '', best_text)):
                best_text = tess_text
        except Exception as e:
            logger.error(f"🔍 {label}: Tesseract OCR falló: {e!r}")

        # ── Escalón 7: GPT-4o Vision OCR (red, paga, último recurso) ──────────
        if vision_fn is not None:
            try:
                logger.error(
                    f"🔍 {label}: intentando GPT-4o Vision OCR (último recurso, paga)..."
                )
                ocr_text = await _ocr_pdf_with_vision(
                    content, vision_fn, file_label=label
                )
                non_ws_ocr = len(re.sub(r'\s+', '', ocr_text or ''))
                if non_ws_ocr >= 50:
                    logger.error(
                        f"🔍 {label}: ✅ Vision OCR rescató el PDF | "
                        f"{len(ocr_text)}c totales, {non_ws_ocr}c no-whitespace"
                    )
                    _cache_put(cache_key, ocr_text)
                    return ocr_text
                logger.error(
                    f"🔍 {label}: Vision OCR también dio poco texto ({non_ws_ocr}c)"
                )
                if non_ws_ocr > len(re.sub(r'\s+', '', best_text)):
                    best_text = ocr_text
            except Exception as e:
                logger.error(f"🔍 {label}: Vision OCR falló: {e!r}")
        else:
            logger.error(
                f"🔍 {label}: vision_fn no fue provista, saltando Vision OCR. "
                f"Letter flows DEBERÍAN pasar call_openai_gpt4o_vision."
            )

        # Devolvemos el mejor intento (probablemente vacío). El caller hace
        # raise ValueError("No se pudo extraer texto del CV...") downstream.
        final_non_ws = len(re.sub(r'\s+', '', best_text))
        logger.error(
            f"📄 {label}: 🚫 TODOS los métodos (5 PDF + Tesseract + Vision) "
            f"agotados. Devolviendo {final_non_ws}c. El caller va a hacer raise."
        )
        return best_text

    if fname.endswith('.docx'):
        try:
            import docx as python_docx
            doc = python_docx.Document(io.BytesIO(content))
            text = "\n".join(p.text for p in doc.paragraphs)
            _cache_put(cache_key, text)
            return text
        except Exception as e:
            logger.warning(f"📄 {label}: python-docx falló: {e!r}")

    text = content.decode('utf-8', errors='ignore')
    _cache_put(cache_key, text)
    return text


async def extract_text_from_upload(
    upload_file,
    *,
    vision_fn: Optional[VisionFn] = None,
    file_label: str = "",
) -> str:
    """Read a FastAPI UploadFile and run it through extract_text_from_bytes."""
    content = await upload_file.read()
    return await extract_text_from_bytes(
        content,
        upload_file.filename or "",
        vision_fn=vision_fn,
        content_type=getattr(upload_file, 'content_type', '') or '',
        file_label=file_label,
    )


# ──────────────────────────────────────────────────────────────────────────────
# Anti-hallucination guard — verify the LLM-extracted signer is in the CV
# ──────────────────────────────────────────────────────────────────────────────

# Generic names GPT-4o falls back to when it has no real data to anchor on.
_DEMO_NAMES = frozenset({
    "emily thompson", "john smith", "jane doe", "sarah johnson",
    "michael chen", "david brown", "maria garcia", "robert davis",
    "james wilson", "patricia martinez", "linda anderson", "barbara taylor",
})


def validate_signer_name_in_cv(
    signer_name: str,
    cv_text: str,
    *,
    role_label: str = "firmante",
) -> None:
    """Verify the LLM-extracted signer name actually appears in the CV.

    Raises ValueError with a user-friendly message if:
      - The name is empty or N/A.
      - The name matches a known demo/hallucination placeholder.
      - None of the ≥3-char name tokens appear in the CV text.

    `role_label` is interpolated into error messages: "firmante",
    "recomendador", "candidato", etc.
    """
    name = (signer_name or "").strip()
    if not name or name.upper() in {"N/A", "NA", "NONE", "UNKNOWN", "TBD"}:
        raise ValueError(
            f"El análisis del CV del {role_label} no devolvió un nombre. "
            f"Verificá que el PDF tenga texto legible y que el nombre esté "
            f"visible al inicio del documento."
        )

    name_lower = name.lower()
    if name_lower in _DEMO_NAMES:
        raise ValueError(
            f"El {role_label} extraído '{name}' es un nombre de demostración "
            f"típico de alucinación de LLM. Revisá el CV adjunto."
        )

    cv_lower = (cv_text or "").lower()
    name_tokens = [t for t in re.split(r'\s+', name_lower) if len(t) >= 3]
    if not name_tokens or not any(t in cv_lower for t in name_tokens):
        logger.error(
            f"🚨 {role_label.capitalize()} alucinado: {name!r} NO aparece en el "
            f"CV. Primeros 500c del CV: {cv_text[:500]!r}"
        )
        raise ValueError(
            f"El nombre del {role_label} extraído ('{name}') no aparece en el "
            f"CV adjunto. Esto suele indicar que el LLM alucinó datos. "
            f"Revisá que el CV cargado corresponda a esa persona."
        )
