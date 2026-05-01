"""File upload handler and text extraction"""

import io
import logging
import pdfplumber
import docx
from fastapi import HTTPException

logger = logging.getLogger(__name__)

def extract_text_from_file(file_content: bytes, filename: str) -> str:
    """
    Extract text from uploaded document - FAST extraction without AI
    Supports: .pdf, .docx, .doc, .txt, .png, .jpg, .jpeg (OCR)
    
    Args:
        file_content: File bytes
        filename: Original filename
        
    Returns:
        Extracted text as string
    """
    try:
        filename_lower = filename.lower()
        
        if filename_lower.endswith(('.docx', '.doc')):
            return _extract_from_docx(file_content)
        elif filename_lower.endswith('.pdf'):
            return _extract_from_pdf(file_content)
        elif filename_lower.endswith('.txt'):
            return _extract_from_txt(file_content)
        elif filename_lower.endswith(('.png', '.jpg', '.jpeg', '.webp', '.heic')):
            return _extract_from_image(file_content, filename)
        else:
            # Return empty string for unsupported formats instead of raising error
            logger.warning(f"Formato no soportado para extracción de texto: {filename}")
            return f"[Archivo: {filename} - formato no soportado para extracción de texto]"
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error extracting text from {filename}: {e}")
        return f"[Error al procesar archivo {filename}: {str(e)}]"

def _extract_from_pdf(file_content: bytes) -> str:
    """Extract text from PDF using multiple methods with fallbacks"""
    text_parts = []
    
    # Método 1: pdfplumber (más rápido y preciso para PDFs con texto)
    try:
        logger.info("Intentando extracción con pdfplumber...")
        with pdfplumber.open(io.BytesIO(file_content)) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text and text.strip():
                    text_parts.append(text)
        
        if text_parts and len('\n'.join(text_parts)) > 100:
            logger.info(f"✅ pdfplumber extrajo {len(''.join(text_parts))} caracteres")
            return '\n\n'.join(text_parts)
    except Exception as e:
        logger.warning(f"pdfplumber falló: {e}")
    
    # Método 2: PyPDF2 (alternativa)
    try:
        logger.info("Intentando extracción con PyPDF2...")
        import PyPDF2
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
        text_parts = []
        for page in pdf_reader.pages:
            text = page.extract_text()
            if text and text.strip():
                text_parts.append(text)
        
        if text_parts and len('\n'.join(text_parts)) > 100:
            logger.info(f"✅ PyPDF2 extrajo {len(''.join(text_parts))} caracteres")
            return '\n\n'.join(text_parts)
    except Exception as e:
        logger.warning(f"PyPDF2 falló: {e}")
    
    # Método 3: PyMuPDF/fitz (muy bueno para PDFs complejos)
    try:
        logger.info("Intentando extracción con PyMuPDF...")
        import fitz  # PyMuPDF
        doc = fitz.open(stream=file_content, filetype="pdf")
        text_parts = []
        for page in doc:
            text = page.get_text()
            if text and text.strip():
                text_parts.append(text)
        doc.close()
        
        if text_parts and len('\n'.join(text_parts)) > 100:
            logger.info(f"✅ PyMuPDF extrajo {len(''.join(text_parts))} caracteres")
            return '\n\n'.join(text_parts)
    except Exception as e:
        logger.warning(f"PyMuPDF falló: {e}")
    
    # Método 4: OCR con pdf2image + pytesseract (para PDFs escaneados)
    try:
        logger.info("Intentando extracción con OCR (pdf2image + pytesseract)...")
        from pdf2image import convert_from_bytes
        import pytesseract
        
        # Convertir PDF a imágenes
        images = convert_from_bytes(file_content, dpi=200, first_page=1, last_page=10)
        
        text_parts = []
        for i, image in enumerate(images):
            text = pytesseract.image_to_string(image, lang='spa+eng')
            if text and text.strip():
                text_parts.append(text)
                logger.info(f"  Página {i+1}: {len(text)} caracteres extraídos")
        
        if text_parts and len('\n'.join(text_parts)) > 100:
            logger.info(f"✅ OCR extrajo {len(''.join(text_parts))} caracteres de {len(images)} páginas")
            return '\n\n'.join(text_parts)
    except ImportError as ie:
        logger.warning(f"OCR no disponible (librería faltante): {ie}")
    except Exception as e:
        logger.warning(f"OCR falló: {e}")
    
    # Método 5: Fallback final - GPT-4o Vision
    try:
        logger.info("Intentando extracción con GPT-4o Vision (fallback final)...")
        text = _extract_with_gpt4o_vision(file_content)
        if text and len(text) > 100:
            logger.info(f"✅ GPT-4o Vision extrajo {len(text)} caracteres")
            return text
    except Exception as e:
        logger.warning(f"GPT-4o Vision falló: {e}")
    
    # Si todos los métodos fallaron
    raise ValueError("No se pudo extraer texto del PDF. El archivo puede estar corrupto, vacío o ser una imagen sin texto reconocible.")

def _extract_with_gpt4o_vision(file_content: bytes) -> str:
    """Extract text from PDF using GPT-4o Vision as last resort"""
    import base64
    import os
    
    try:
        from openai import OpenAI
        from pdf2image import convert_from_bytes
        
        images = convert_from_bytes(file_content, dpi=150, first_page=1, last_page=5)
        
        if not images:
            raise ValueError("No se pudo convertir PDF a imagen")
        
        client = OpenAI(api_key=os.environ.get('OPENAI_API_KEY'))
        
        all_text = []
        for i, image in enumerate(images[:5]):
            img_buffer = io.BytesIO()
            image.save(img_buffer, format='PNG')
            img_base64 = base64.b64encode(img_buffer.getvalue()).decode('utf-8')
            
            response = client.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {
                        "role": "user",
                        "content": [
                            {
                                "type": "text",
                                "text": "Extract ALL text from this document image. Preserve the structure (headings, paragraphs, lists). Return ONLY the extracted text, no explanations."
                            },
                            {
                                "type": "image_url",
                                "image_url": {
                                    "url": f"data:image/png;base64,{img_base64}",
                                    "detail": "high"
                                }
                            }
                        ]
                    }
                ],
                max_tokens=4000
            )
            
            page_text = response.choices[0].message.content
            if page_text:
                all_text.append(f"--- Página {i+1} ---\n{page_text}")
        
        return '\n\n'.join(all_text)
        
    except Exception as e:
        logger.error(f"GPT-4o Vision extraction failed: {e}")
        raise

def _extract_from_image(file_content: bytes, filename: str) -> str:
    """Extract text from image using OCR (pytesseract)"""
    try:
        logger.info(f"Intentando OCR para imagen: {filename}")
        
        # Método 1: pytesseract OCR
        try:
            import pytesseract
            from PIL import Image
            
            image = Image.open(io.BytesIO(file_content))
            if image.mode != 'RGB':
                image = image.convert('RGB')
            
            text = pytesseract.image_to_string(image, lang='spa+eng')
            
            if text and text.strip() and len(text.strip()) > 20:
                logger.info(f"✅ OCR extrajo {len(text)} caracteres de {filename}")
                return text.strip()
        except ImportError:
            logger.warning("pytesseract no disponible, intentando método alternativo")
        except Exception as e:
            logger.warning(f"OCR con pytesseract falló: {e}")
        
        # Si OCR falla, devolver descripción del archivo
        logger.warning(f"No se pudo extraer texto de la imagen {filename}")
        return f"[Imagen: {filename} - Se requiere revisión manual. El contenido de la imagen no pudo ser extraído automáticamente.]"
        
    except Exception as e:
        logger.error(f"Error in image extraction for {filename}: {e}")
        return f"[Error procesando imagen {filename}: {str(e)}]"

def _extract_from_docx(file_content: bytes) -> str:
    """Extract text from DOCX (including tables, and images via GPT-4o Vision)"""
    try:
        doc = docx.Document(io.BytesIO(file_content))
        text_parts = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        text_parts.append(cell_text)
        
        # If text was found, return it
        if text_parts and len('\n\n'.join(text_parts)) > 50:
            return '\n\n'.join(text_parts)
        
        # No text found — try to extract embedded images and OCR them with GPT-4o Vision
        logger.warning("DOCX tiene poco o ningún texto. Buscando imágenes embebidas para OCR...")
        image_texts = _extract_docx_images_with_vision(file_content)
        if image_texts:
            logger.info(f"✅ GPT-4o Vision extrajo texto de {len(image_texts)} imagen(es) del DOCX")
            return '\n\n'.join(image_texts)
        
        raise ValueError("El documento DOCX no contiene texto ni imágenes reconocibles")
    except ValueError:
        raise
    except Exception as e:
        logger.error(f"Error in DOCX extraction: {e}")
        raise ValueError(f"Error al leer DOCX: {str(e)}")


def _extract_docx_images_with_vision(file_content: bytes) -> list:
    """Extract images from a DOCX file and OCR them with GPT-4o Vision."""
    import zipfile
    import base64
    import os
    
    image_texts = []
    
    try:
        # DOCX files are ZIP archives; images are in word/media/
        with zipfile.ZipFile(io.BytesIO(file_content)) as z:
            image_files = [
                name for name in z.namelist()
                if name.startswith('word/media/') and
                any(name.lower().endswith(ext) for ext in ('.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tiff', '.webp'))
            ]
        
        if not image_files:
            logger.warning("No se encontraron imágenes embebidas en el DOCX")
            return []
        
        logger.info(f"Encontradas {len(image_files)} imagen(es) en DOCX. Procesando con GPT-4o Vision...")
        
        # Use OpenRouter (same pattern as rest of the app)
        import httpx
        openrouter_key = os.environ.get('OPENROUTER_API_KEY', '')
        
        with zipfile.ZipFile(io.BytesIO(file_content)) as z:
            for img_name in image_files[:5]:  # Max 5 images
                try:
                    img_bytes = z.read(img_name)
                    ext = img_name.split('.')[-1].lower()
                    mime = f"image/{'jpeg' if ext in ('jpg', 'jpeg') else ext}"
                    img_b64 = base64.b64encode(img_bytes).decode('utf-8')
                    
                    import asyncio
                    import concurrent.futures
                    
                    async def call_vision(b64, mime_type):
                        async with httpx.AsyncClient(timeout=60.0) as client:
                            resp = await client.post(
                                "https://openrouter.ai/api/v1/chat/completions",
                                headers={
                                    "Authorization": f"Bearer {openrouter_key}",
                                    "Content-Type": "application/json",
                                    "HTTP-Referer": "https://redaccion.urpeintegralservices.co",
                                },
                                json={
                                    "model": "openai/gpt-4o",
                                    "messages": [{
                                        "role": "user",
                                        "content": [
                                            {
                                                "type": "text",
                                                "text": "Extract ALL text from this document image. Preserve structure (headings, paragraphs, lists). Return ONLY the extracted text, no explanations."
                                            },
                                            {
                                                "type": "image_url",
                                                "image_url": {
                                                    "url": f"data:{mime_type};base64,{b64}",
                                                    "detail": "high"
                                                }
                                            }
                                        ]
                                    }],
                                    "max_tokens": 4000
                                }
                            )
                        if resp.status_code == 200:
                            return resp.json()["choices"][0]["message"]["content"]
                        return None
                    
                    # Run async call - handle both sync and async contexts
                    try:
                        loop = asyncio.get_event_loop()
                        if loop.is_running():
                            # Running from an async context (FastAPI) - use thread executor
                            with concurrent.futures.ThreadPoolExecutor() as executor:
                                future = executor.submit(asyncio.run, call_vision(img_b64, mime))
                                text = future.result(timeout=60)
                        else:
                            text = loop.run_until_complete(call_vision(img_b64, mime))
                    except RuntimeError:
                        # Fallback: create new event loop in thread
                        with concurrent.futures.ThreadPoolExecutor() as executor:
                            future = executor.submit(asyncio.run, call_vision(img_b64, mime))
                            text = future.result(timeout=60)
                    
                    if text and text.strip():
                        image_texts.append(text.strip())
                        logger.info(f"  ✅ Imagen {img_name}: {len(text)} chars extraídos")
                        
                except Exception as img_err:
                    logger.warning(f"  ⚠️ No se pudo procesar imagen {img_name}: {img_err}")
    
    except Exception as e:
        logger.error(f"Error extracting DOCX images: {e}")
    
    return image_texts

def _extract_from_txt(file_content: bytes) -> str:
    """Extract text from TXT"""
    try:
        text = file_content.decode('utf-8')
        if not text.strip():
            raise ValueError("El archivo TXT está vacío")
        return text
    except UnicodeDecodeError:
        try:
            text = file_content.decode('latin-1')
            if not text.strip():
                raise ValueError("El archivo TXT está vacío")
            return text
        except Exception as e:
            logger.error(f"Error in TXT extraction: {e}")
            raise ValueError(f"Error al leer TXT: {str(e)}")
