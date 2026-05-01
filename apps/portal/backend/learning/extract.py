"""Text extraction from PDF and DOCX files."""

import io
import logging

logger = logging.getLogger(__name__)


def extract_text(file_bytes: bytes, mime_type: str, filename: str) -> str:
    """Extract plain text from PDF or DOCX bytes."""
    name_lower = (filename or "").lower()

    is_pdf = mime_type == "application/pdf" or name_lower.endswith(".pdf")
    is_docx = (
        mime_type
        in (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "application/msword",
        )
        or name_lower.endswith(".docx")
    )

    if is_pdf:
        return _extract_pdf(file_bytes)
    if is_docx:
        return _extract_docx(file_bytes)
    raise ValueError(f"Unsupported file type: {mime_type or filename}")


def _extract_pdf(data: bytes) -> str:
    import fitz  # PyMuPDF

    text_parts = []
    with fitz.open(stream=data, filetype="pdf") as doc:
        for page_num, page in enumerate(doc, start=1):
            page_text = page.get_text("text") or ""
            if page_text.strip():
                text_parts.append(f"[Página {page_num}]\n{page_text.strip()}")
    return "\n\n".join(text_parts)


def _extract_docx(data: bytes) -> str:
    from docx import Document

    doc = Document(io.BytesIO(data))
    parts = []
    for para in doc.paragraphs:
        if para.text and para.text.strip():
            parts.append(para.text.strip())
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)
